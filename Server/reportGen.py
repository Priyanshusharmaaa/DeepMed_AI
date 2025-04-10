import os
from PIL import Image as PILImage
from agno.agent import Agent
from agno.models.google import Gemini
from agno.tools.duckduckgo import DuckDuckGoTools
from agno.media import Image as AgnoImage
import streamlit as st
import re
from phrase_ground import plot_phrase_grounding, save_image_and_prompt_to_docx
from pathlib import Path
from docx import Document
from docx.shared import Inches
import io


GOOGLE_API_KEY = "AIzaSyB7UNr5KkfAoMgFlnS9ed1v6ZmhehrrHq0"
os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY


if not GOOGLE_API_KEY:
    raise ValueError("⚠️ Please set your Google API Key in GOOGLE_API_KEY")


medical_agent = Agent(
    model=Gemini(id="gemini-2.0-flash-exp"),
    tools=[DuckDuckGoTools()],
    markdown=True
)


query = """
You are a highly skilled medical imaging expert with extensive knowledge in radiology and diagnostic imaging. Analyze the medical image and structure your response as follows:

### 1. Image Type & Region
- Identify imaging modality (X-ray/MRI/CT/Ultrasound/etc.).
- Specify anatomical region and positioning.
- Evaluate image quality and technical adequacy.

### 2. Possible Category
- Amongst the 14 categories : Atelectasis, Cardiomegaly, Effusion, Infiltration, Mass, Nodule, Pneumonia, Pneumothorax, Consolidation, Edema, Emphysema, Fibrosis, Pleural Thickening, Hernia, Normal.

### 3. Key Findings
- Highlight primary observations systematically.
- Identify potential abnormalities with detailed descriptions.
- Include measurements and densities where relevant.

### 4. Diagnostic Assessment
- Provide primary diagnosis with confidence level.
- List differential diagnoses ranked by likelihood.
- Support each diagnosis with observed evidence.
- Highlight critical/urgent findings.

### 5. Patient-Friendly Explanation
- Simplify findings in clear, non-technical language.
- Avoid medical jargon or provide easy definitions.
- Include relatable visual analogies.

### 6. Research Context
- Use DuckDuckGo search to find recent medical literature.
- Provide 2-3 key references supporting the analysis.
- Conclude the report

Ensure a structured and medically accurate response using clear markdown formatting.
"""

def extract_key_findings(report_text: str) -> list:
    """
    Extracts the 'Key Findings' section from the AI medical report.

    Returns:
        List of extracted findings as phrases/sentences.
    """
    key_findings_match = re.search(r"### 3\. Key Findings\n(.+?)\n###", report_text, re.DOTALL)
    
    if key_findings_match:
        findings_block = key_findings_match.group(1)
        findings = [
            line.strip("-• ").strip()
            for line in findings_block.strip().split("\n")
            if line.strip() and not line.lower().startswith("note:")
        ]
        return findings
    else:
        return []

def analyze_medical_image(image_path):
    """Processes and analyzes a medical image using AI."""
    image = PILImage.open(image_path)
    width, height = image.size
    aspect_ratio = width / height
    new_width = 500
    new_height = int(new_width / aspect_ratio)
    resized_image = image.resize((new_width, new_height))

    temp_path = "temp_resized_image.png"
    resized_image.save(temp_path)

    agno_image = AgnoImage(filepath=temp_path)

    try:
        response = medical_agent.run(query, images=[agno_image])
        return response.content
    except Exception as e:
        return f"⚠️ Analysis error: {e}"
    finally:
        os.remove(temp_path)


st.set_page_config(page_title="Medical Image Analysis", layout="centered",page_icon="🩺")
st.title("🩺 Medical Image Analysis Tool 🔬")
st.markdown(
    """
    Welcome to the **Medical Image Analysis** tool! 📸  
    Upload a medical image (X-ray, MRI, CT, Ultrasound, etc.), and our AI-powered system will analyze it, providing detailed findings, diagnosis, and research insights.  
    Let's get started!
    """
)


st.sidebar.header("Upload Your Medical Image:")
uploaded_file = st.sidebar.file_uploader("Choose a medical image file", type=["jpg", "jpeg", "png", "bmp", "gif"])

if uploaded_file is not None:
    st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)
    
    if st.sidebar.button("Analyze Image"):
        with st.spinner("🔍 Analyzing the image... Please wait."):
            image_ext = uploaded_file.type.split('/')[1]
            image_path = f"temp_image.{image_ext}"
            with open(image_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            report = analyze_medical_image(image_path)
            image_path = Path(image_path)
            doc = Document()
            doc.add_heading("🩺 Medical Image Analysis Report", level=1)

            for line in report.split("\n"):
                line = line.strip()
                if line.startswith("###"):
                    doc.add_heading(line.replace("###", "").strip(), level=2)
                elif line.startswith("**") and line.endswith("**"):
                    doc.add_heading(line.strip("**"), level=3)
                elif line:
                    doc.add_paragraph(line)

            
            findings = extract_key_findings(report)
            for finding in findings:
                result = plot_phrase_grounding(image_path, finding)
                result.save("grounded.png")
                save_image_and_prompt_to_docx(Path("grounded.png"), finding, doc)

            
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

           
            st.subheader("📋 Analysis Report")
            st.markdown(report, unsafe_allow_html=True)

            st.download_button(
                label="📥 Download Full Report (DOCX)",
                data=docx_buffer,
                file_name="medical_analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            os.remove(image_path)
            
else:
    st.warning("⚠️ Please upload a medical image to begin analysis.")
