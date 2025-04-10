from pathlib import Path
from typing import List, Tuple
import requests
from PIL import Image
from io import BytesIO
import matplotlib.pyplot as plt
import torch
from IPython.display import display, Markdown
import io
from health_multimodal.common.visualization import plot_phrase_grounding_similarity_map
from health_multimodal.text import get_bert_inference
from health_multimodal.text.utils import BertEncoderType
from health_multimodal.image import get_image_inference
from health_multimodal.image.utils import ImageModelType
from health_multimodal.vlp import ImageTextInferenceEngine
from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading("Image segmentation results", level=1)


text_inference = get_bert_inference(BertEncoderType.BIOVIL_T_BERT)
image_inference = get_image_inference(ImageModelType.BIOVIL_T)
image_text_inference = ImageTextInferenceEngine(
    image_inference_engine=image_inference,
    text_inference_engine=text_inference,
)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
image_text_inference.to(device)

TypeBox = Tuple[float, float, float, float]

def save_image_and_prompt_to_docx(img_path, text_prompt: str,doc : Document) -> None:
    

  
    doc.add_picture(str(img_path), width=Inches(6))  
   
    doc.add_paragraph(f"{text_prompt}", style='Normal')


   
    

    print(f"Images added")

def plot_phrase_grounding(image_path: Path, text_prompt: str) -> None:
    similarity_map = image_text_inference.get_similarity_map_from_raw_data(
        image_path=image_path,
        query_text=text_prompt,
        interpolation="bilinear",
    )
    
    fig = plot_phrase_grounding_similarity_map(
        image_path=image_path,
        similarity_map=similarity_map,
    )

    
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)  
    buf.seek(0)

    
    pil_img = Image.open(buf).convert("RGB")
    return pil_img
text_prompt = ["Hyperlucency in the right lung field","Tracheal deviation towards the left",
               "Elevated left hemidiaphragm","Loss of volume in the left lung with mediastinal shift"]
bboxes = [
    (306, 168, 124, 101),
]

text = (
    'The ground-truth bounding box annotation for the phrase'
    f' *{text_prompt}* is shown in the middle figure (in black).'
)

